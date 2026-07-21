"""Resume text extraction for PDF, DOCX, and TXT.

Pure extraction — no AI, no interpretation. We store exactly what the file contains so
later phases never invent facts about the candidate.
"""

from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


class UnsupportedResumeType(ValueError):
    """Raised for a file that isn't PDF, DOCX, or TXT."""


def _extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts).strip()


def _parse_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Include table cell text too — many resumes lay content out in tables.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part).strip()


def _parse_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


def parse_resume(filename: str, data: bytes) -> str:
    """Extract plain text from a resume file, dispatched by extension."""
    ext = _extension(filename)
    if ext == "pdf":
        return _parse_pdf(data)
    if ext == "docx":
        return _parse_docx(data)
    if ext == "txt":
        return _parse_txt(data)
    raise UnsupportedResumeType(
        f"Unsupported file type '.{ext}'. Allowed: PDF, DOCX, TXT."
    )
