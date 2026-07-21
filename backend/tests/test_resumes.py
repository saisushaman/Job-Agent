"""Tests for the resume system: parser + upload/version/preview endpoints."""

from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient
from fpdf import FPDF

from app.services.resume_parser import UnsupportedResumeType, parse_resume


# --- parser unit tests ------------------------------------------------------


def test_parse_txt() -> None:
    text = parse_resume("resume.txt", b"Jane Doe\nPython, FastAPI")
    assert "Jane Doe" in text
    assert "FastAPI" in text


def test_parse_docx() -> None:
    doc = Document()
    doc.add_paragraph("John Smith")
    doc.add_paragraph("AI Engineer with local LLM experience")
    buf = io.BytesIO()
    doc.save(buf)
    text = parse_resume("resume.docx", buf.getvalue())
    assert "John Smith" in text
    assert "AI Engineer" in text


def test_parse_pdf() -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 10, "Cloud Engineer Resume")
    data = bytes(pdf.output())
    text = parse_resume("resume.pdf", data)
    assert "Cloud Engineer" in text


def test_parse_unsupported() -> None:
    try:
        parse_resume("resume.rtf", b"data")
        assert False, "expected UnsupportedResumeType"
    except UnsupportedResumeType:
        pass


# --- endpoint tests ---------------------------------------------------------


def test_default_tracks_seeded(client: TestClient) -> None:
    resp = client.get("/api/resumes")
    assert resp.status_code == 200
    names = {r["name"] for r in resp.json()}
    assert {
        "Master Resume",
        "Software Engineer",
        "AI Engineer",
        "Cloud Engineer",
        "DevOps Engineer",
    } <= names


def test_upload_and_version_flow(client: TestClient) -> None:
    resume_id = client.get("/api/resumes").json()[0]["id"]

    # First upload -> version 1
    r1 = client.post(
        f"/api/resumes/{resume_id}/versions",
        files={"file": ("cv.txt", b"Alice Example\nSkills: Python", "text/plain")},
    )
    assert r1.status_code == 201, r1.text
    v1 = r1.json()
    assert v1["version_number"] == 1
    assert "Alice Example" in v1["parsed_text"]

    # Second upload -> version 2
    r2 = client.post(
        f"/api/resumes/{resume_id}/versions",
        files={"file": ("cv2.txt", b"Alice Example v2", "text/plain")},
    )
    assert r2.json()["version_number"] == 2

    # Track detail shows version history
    detail = client.get(f"/api/resumes/{resume_id}").json()
    assert len(detail["versions"]) == 2

    # List shows latest version number
    summary = next(x for x in client.get("/api/resumes").json() if x["id"] == resume_id)
    assert summary["latest_version_number"] == 2
    assert summary["version_count"] == 2

    # Preview parsed text
    preview = client.get(f"/api/resumes/versions/{v1['id']}").json()
    assert "Alice Example" in preview["parsed_text"]


def test_upload_rejects_unsupported_type(client: TestClient) -> None:
    resume_id = client.get("/api/resumes").json()[0]["id"]
    resp = client.post(
        f"/api/resumes/{resume_id}/versions",
        files={"file": ("cv.rtf", b"data", "application/rtf")},
    )
    assert resp.status_code == 400


def test_upload_unknown_track_404(client: TestClient) -> None:
    resp = client.post(
        "/api/resumes/999999/versions",
        files={"file": ("cv.txt", b"data", "text/plain")},
    )
    assert resp.status_code == 404
